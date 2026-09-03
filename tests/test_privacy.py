"""Privacy-focus verification suite.

Covers the deeper privacy guarantees beyond basic metadata stripping:
  * GPS location detection & human-readable reporting (JPEG / PNG eXIf /
    TIFF / HEIC), with the coordinates actually removed from the output
  * PDF document ID (/ID) fingerprint reset
  * PDF /PieceInfo editing-history removal (incl. page-level XMP)
  * WordprocessingML edit-session ID (w:rsid) stripping
  * Macro-reference scrubbing in [Content_Types].xml and *.rels
  * end-to-end "output has no leftover identity markers" sweeps

Run:  python tests/test_privacy.py
"""
import os
import shutil
import sys
import zipfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import scrubber
from scrubber import (sanitize_file, Options, default_output_path,
                      _gps_from_exif, _strip_rsids, _strip_vba_refs)

from PIL import Image
from PIL.TiffImagePlugin import IFDRational

TMP = os.path.join(os.path.dirname(os.path.abspath(__file__)), "tmp_priv")
OUT = os.path.join(TMP, "out")
PASS = FAIL = 0


def check(name, cond, detail=""):
    global PASS, FAIL
    if cond:
        PASS += 1
        print(f"  PASS  {name}")
    else:
        FAIL += 1
        print(f"  FAIL  {name}   {detail}")


def sanitize(src, opts=None, out_dir=OUT):
    opts = opts or Options()
    dst = default_output_path(src, out_dir, opts.suffix)
    return sanitize_file(src, dst, opts)


def make_gps_exif():
    exif = Image.Exif()
    exif[0x010F] = "Acme Corp"
    exif[0x0110] = "SecretCam 9000"
    exif[0x8825] = {1: "S",
                    2: (IFDRational(31, 1), IFDRational(57, 1), IFDRational(1, 1)),
                    3: "E",
                    4: (IFDRational(115, 1), IFDRational(51, 1), IFDRational(1, 1))}
    return exif


# ------------------------------------------------------------- GPS reporting

def t_gps_unit():
    print("\n[GPS parser (unit)]")
    exif = make_gps_exif().tobytes()
    s = _gps_from_exif(exif)
    check("extracts coordinates", s is not None and "31" in s and "115" in s,
          str(s))
    check("has hemisphere letters", "S" in s and "E" in s, str(s))
    check("no prefix variant", _gps_from_exif(exif[6:]) == s)
    # no-GPS EXIF -> None
    e2 = Image.Exif(); e2[0x010F] = "Acme"
    check("no GPS -> None", _gps_from_exif(e2.tobytes()) is None)
    # hostile inputs never raise
    for junk in (b"", b"\x00", b"Exif\x00\x00", b"MM\x00*\x00\x00\x00\x08",
                 b"\xff" * 100, b"II*\x00\xff\xff\xff\xff"):
        try:
            _gps_from_exif(junk)
            check(f"junk {junk[:8]!r} handled", True)
        except Exception as e:
            check(f"junk {junk[:8]!r} handled", False, repr(e))


def t_gps_jpeg():
    print("\n[GPS detection: JPEG]")
    src = os.path.join(TMP, "gps.jpg")
    Image.new("RGB", (64, 48), (200, 40, 40)).save(
        src, "JPEG", exif=make_gps_exif().tobytes())
    rep = sanitize(src)
    check("reports GPS location", "GPS location" in rep.summary(), rep.summary())
    check("reports coordinates", "31" in rep.summary() and "115" in rep.summary(),
          rep.summary())
    check("EXIF also reported", "EXIF" in rep.summary())
    with Image.open(rep.dst) as im:
        check("GPS actually removed", not dict(im.getexif()))
    # exif WITHOUT gps does not claim GPS
    src2 = os.path.join(TMP, "nogps.jpg")
    e = Image.Exif(); e[0x010F] = "Acme"
    Image.new("RGB", (32, 32), (1, 2, 3)).save(src2, "JPEG", exif=e.tobytes())
    rep2 = sanitize(src2)
    check("no false GPS report", "GPS" not in rep2.summary(), rep2.summary())


def t_gps_png():
    print("\n[GPS detection: PNG eXIf]")
    src = os.path.join(TMP, "gps.png")
    Image.new("RGB", (32, 32), (10, 120, 200)).save(
        src, "PNG", exif=make_gps_exif().tobytes())
    rep = sanitize(src)
    check("PNG reports GPS", "GPS location" in rep.summary(), rep.summary())
    # verify output PNG has no eXIf chunk at all
    import struct
    data = open(rep.dst, "rb").read()
    i, n = 8, len(data)
    chunks = []
    while i + 8 <= n:
        length = struct.unpack(">I", data[i:i + 4])[0]
        chunks.append(data[i + 4:i + 8])
        i += 12 + length
    check("eXIf chunk gone", b"eXIf" not in chunks, str(chunks))


def t_gps_tiff_heic():
    print("\n[GPS detection: TIFF & HEIC re-encode path]")
    # TIFF
    src = os.path.join(TMP, "gps.tiff")
    img = Image.new("RGB", (40, 30), (11, 22, 33))
    try:
        img.save(src, "TIFF", exif=make_gps_exif().tobytes())
    except Exception:
        img.save(src, "TIFF")
    rep = sanitize(src)
    check("TIFF reports GPS", "GPS location" in rep.summary(), rep.summary())
    with Image.open(rep.dst) as im:
        exif = dict(im.getexif())
        check("TIFF GPS removed",
              not any(t in exif for t in (271, 272, 305, 306, 34853)))

    # HEIC (if available)
    if not scrubber.HAS_HEIF:
        print("  SKIP  HEIC not available")
        return
    h = os.path.join(TMP, "gps.heic")
    try:
        Image.new("RGB", (48, 36), (70, 80, 90)).save(
            h, format="HEIF", exif=make_gps_exif().tobytes())
        src_has = bool(dict(Image.open(h).getexif()))
    except Exception as e:
        Image.new("RGB", (48, 36), (70, 80, 90)).save(h, format="HEIF")
        src_has = False
        print(f"  note: HEIC exif-save unsupported ({e}); strip still tested")
    rep = sanitize(h)
    check("HEIC processed", rep.ok, rep.summary())
    with Image.open(rep.dst) as im:
        check("HEIC EXIF gone", not dict(im.getexif()))
    if src_has:
        check("HEIC reports GPS", "GPS location" in rep.summary(), rep.summary())


# ------------------------------------------------------------- PDF privacy

def t_pdf_id_pieceinfo():
    print("\n[PDF: /ID fingerprint + /PieceInfo]")
    import pikepdf
    src = os.path.join(TMP, "doc.pdf")
    pdf = pikepdf.new()
    pdf.add_blank_page(page_size=(200, 200))
    pdf.docinfo["/Author"] = "Alice"
    pdf.Root["/PieceInfo"] = pikepdf.Dictionary({
        "/Illustrator": pikepdf.Dictionary({
            "/LastModified": pikepdf.String("D:20240101000000Z")})})
    pdf.save(src)
    # NOTE: always keep pikepdf Pdf objects in a variable — chaining
    # pikepdf.open(...).trailer[...] as a temporary can crash pikepdf.
    src_pdf = pikepdf.open(src)
    id_before = str(src_pdf.trailer["/ID"])
    src_pdf.close()

    rep = sanitize(src)
    check("pdf cleaned", rep.status == "ok", rep.summary())
    check("ID reset reported", "document ID" in rep.summary(), rep.summary())
    check("PieceInfo reported", "PieceInfo" in rep.summary()
          or "metadata object" in rep.summary(), rep.summary())
    q = pikepdf.open(rep.dst)
    id_after = str(q.trailer["/ID"])
    check("ID fingerprint changed", id_before != id_after)
    check("ID first element zeroed",
          str(q.trailer["/ID"][0]) == "\x00" * 16,
          repr(str(q.trailer["/ID"][0])))
    check("PieceInfo gone", "/PieceInfo" not in q.Root)
    check("author gone", not list(q.docinfo.keys()))
    q.close()


def t_pdf_page_xmp():
    print("\n[PDF: page-level XMP removal]")
    import pikepdf
    src = os.path.join(TMP, "pagexmp.pdf")
    pdf = pikepdf.new()
    pdf.add_blank_page(page_size=(200, 200))
    # page-level /Metadata stream (XMP can hide anywhere)
    pdf.pages[0].obj["/Metadata"] = pdf.make_stream(b"<x:xmpmeta/>")
    pdf.save(src)
    rep = sanitize(src)
    check("pdf cleaned", rep.status == "ok", rep.summary())
    q = pikepdf.open(rep.dst)
    check("page-level XMP removed", "/Metadata" not in q.pages[0].obj)
    q.close()


# ------------------------------------------------------------- Office privacy

def t_rsid_stripping():
    print("\n[WordprocessingML rsid edit-session IDs]")
    # unit
    xml = (b'<w:document xmlns:w="w"><w:body><w:p w:rsidR="00AB12CD" '
           b'w:rsidRDefault="00EF34AB"><w:rsids><w:rsid w:val="00AB12CD"/>'
           b'</w:rsids><w:r w:rsidRPr="00DEAD01"><w:t>hi</w:t></w:r></w:p>'
           b'</w:body></w:document>')
    out = _strip_rsids(xml)
    check("rsid attributes removed", b"rsidR" not in out and
          b"rsidRDefault" not in out and b"rsidRPr" not in out, out)
    check("rsids block removed", b"<w:rsids>" not in out)
    check("content kept", b"<w:t>hi</w:t>" in out)

    # end-to-end: docx with rsids
    import docx
    d = os.path.join(TMP, "rsid.docx")
    doc = docx.Document()
    doc.add_paragraph("hello rsid")
    doc.save(d)
    # inject rsid attributes into document.xml
    with zipfile.ZipFile(d) as z:
        entries = {n: z.read(n) for n in z.namelist()}
    docxml = entries["word/document.xml"]
    docxml = docxml.replace(b"<w:body>", b'<w:body w:rsidR="00AB12CD">')
    with zipfile.ZipFile(d, "w") as z:
        for n, data in entries.items():
            z.writestr(n, docxml if n == "word/document.xml" else data)
    rep = sanitize(d)
    check("rsid docx cleaned", rep.status == "ok", rep.summary())
    check("rsids reported", "rsid" in rep.summary(), rep.summary())
    with zipfile.ZipFile(rep.dst) as z:
        docxml_out = z.read("word/document.xml")
    check("rsid gone from output", b"rsidR" not in docxml_out)
    check("doc still opens", "hello rsid" in
          [p.text for p in docx.Document(rep.dst).paragraphs])


def t_macro_ref_scrubbing():
    print("\n[Macro-reference scrubbing in [Content_Types].xml & .rels]")
    # unit
    ct = (b'<?xml version="1.0"?><Types xmlns="x">'
          b'<Override PartName="/word/document.xml" ContentType="ct1"/>'
          b'<Override PartName="/word/vbaProject.bin" ContentType="ct2"/>'
          b'<Override PartName="/word/vbaData.xml" ContentType="ct3"/>'
          b'</Types>')
    out = _strip_vba_refs(ct)
    check("vba Override removed", b"vbaProject" not in out
          and b"vbaData" not in out, out)
    check("other Override kept", b"word/document.xml" in out)

    rels = (b'<?xml version="1.0"?><Relationships xmlns="r">'
            b'<Relationship Id="rId1" Type="t1" Target="document.xml"/>'
            b'<Relationship Id="rId2" Type="t2" Target="vbaProject.bin"/>'
            b'<Relationship Id="rId3" Type="t3" Target="vbaData.xml"/>'
            b'</Relationships>')
    out = _strip_vba_refs(rels)
    check("vba Relationship removed", b"vbaProject" not in out
          and b"vbaData" not in out, out)
    check("other Relationship kept", b"document.xml" in out)

    # end-to-end with realistic macro structure
    import docx
    d = os.path.join(TMP, "macro2.docx")
    doc = docx.Document()
    doc.core_properties.author = "Alice"
    doc.add_paragraph("macro doc")
    doc.save(d)
    # add macro parts + realistic references
    with zipfile.ZipFile(d) as z:
        entries = {n: z.read(n) for n in z.namelist()}
    entries["word/vbaProject.bin"] = b"\xd0\xcf\x11\xe0" + b"M" * 64
    entries["word/vbaData.xml"] = b'<w:vbaSupports xmlns:w="w"/>'
    ct = entries["[Content_Types].xml"]
    ct = ct.replace(b"</Types>",
                    b'<Override PartName="/word/vbaProject.bin" '
                    b'ContentType="application/vnd.ms-office.vbaProject"/>'
                    b'<Override PartName="/word/vbaData.xml" '
                    b'ContentType="application/vnd.ms-word.vbaData+xml"/>'
                    b"</Types>")
    entries["[Content_Types].xml"] = ct
    rels_name = "word/_rels/document.xml.rels"
    rels = entries[rels_name]
    rels = rels.replace(
        b"</Relationships>",
        b'<Relationship Id="rIdMacro" Type="http://schemas.openxmlformats.org/'
        b'officeDocument/2006/relationships/vbaProject" Target="vbaProject.bin"/>'
        b'<Relationship Id="rIdMacro2" Type="http://schemas.openxmlformats.org/'
        b'officeDocument/2006/relationships/vbaData" Target="vbaData.xml"/>'
        b"</Relationships>")
    entries[rels_name] = rels
    with zipfile.ZipFile(d, "w") as z:
        for n, data in entries.items():
            z.writestr(n, data)

    rep = sanitize(d)
    check("macro docx cleaned", rep.status == "ok", rep.summary())
    with zipfile.ZipFile(rep.dst) as z:
        names = z.namelist()
        check("vbaProject.bin removed", "word/vbaProject.bin" not in names)
        check("vbaData.xml removed", "word/vbaData.xml" not in names)
        ct_out = z.read("[Content_Types].xml")
        rels_out = z.read(rels_name)
        check("no vba refs in Content_Types",
              b"vbaProject" not in ct_out and b"vbaData" not in ct_out)
        check("no vba refs in rels",
              b"vbaProject" not in rels_out and b"vbaData" not in rels_out)
    # still opens
    check("docx still opens", "macro doc" in
          [p.text for p in docx.Document(rep.dst).paragraphs])
    check("author removed too", "Alice" not in
          zipfile.ZipFile(rep.dst).read("docProps/core.xml").decode())


def t_sweep_no_identity():
    print("\n[Sweep: output has no leftover identity markers]")
    src = os.path.join(TMP, "sweep.jpg")
    Image.new("RGB", (64, 48), (9, 9, 9)).save(src, "JPEG",
                                               exif=make_gps_exif().tobytes(),
                                               comment=b"me")
    rep = sanitize(src)
    data = open(rep.dst, "rb").read()
    for marker in (b"Exif", b"xmp", b"GPS", b"Acme", b"SecretCam",
                   b"LeakySoft", b"http://ns.adobe.com"):
        check(f"no {marker!r} in output", marker not in data, marker)
    with Image.open(rep.dst) as im:
        check("no exif in decoded image", not dict(im.getexif()))
        check("no info keys leaked", not any(k in im.info for k in
              ("exif", "comment", "xmp")))


def main():
    shutil.rmtree(TMP, ignore_errors=True)
    os.makedirs(OUT)
    t_gps_unit()
    t_gps_jpeg()
    t_gps_png()
    t_gps_tiff_heic()
    t_pdf_id_pieceinfo()
    t_pdf_page_xmp()
    t_rsid_stripping()
    t_macro_ref_scrubbing()
    t_sweep_no_identity()
    print(f"\n===== {PASS} passed, {FAIL} failed =====")
    sys.exit(1 if FAIL else 0)


if __name__ == "__main__":
    main()
