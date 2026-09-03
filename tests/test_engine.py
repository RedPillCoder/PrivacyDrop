"""End-to-end tests for the PrivacyDrop sanitization engine."""
import io
import os
import sys
import shutil
import struct
import zipfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from scrubber import (sanitize_file, Options, strip_jpeg_bytes, strip_png_bytes,
                      strip_webp_bytes, default_output_path)

from PIL import Image

TMP = os.path.join(os.path.dirname(os.path.abspath(__file__)), "tmp")
OUT = os.path.join(TMP, "out")
PASS = FAIL = 0


def check(name, cond, detail=""):
    global PASS, FAIL
    if cond:
        PASS += 1
        print(f"  PASS  {name}")
    else:
        FAIL += 1
        print(f"  FAIL  {name}  {detail}")


def make_jpeg(path):
    """Create a JPEG carrying EXIF (camera + GPS), XMP and a comment."""
    from PIL.TiffImagePlugin import IFDRational
    img = Image.new("RGB", (64, 48), (200, 40, 40))
    exif = Image.Exif()
    exif[0x010F] = "Acme Corp"            # Make
    exif[0x0110] = "SecretCam 9000"       # Model
    exif[0x0131] = "TopSecretSoft 1.0"    # Software
    exif[0x0132] = "2024:01:02 03:04:05"  # DateTime
    exif[0x8825] = {1: "S",
                    2: (IFDRational(31, 1), IFDRational(57, 1), IFDRational(1, 1)),
                    3: "E",
                    4: (IFDRational(115, 1), IFDRational(51, 1), IFDRational(1, 1))}
    img.save(path, "JPEG", quality=92, exif=exif.tobytes(), comment=b"hello")


def read_jpeg_segments(path):
    data = open(path, "rb").read()
    i, n = 2, len(data)
    segs = []
    while i < n - 1:
        if data[i] != 0xFF:
            break
        b = data[i + 1]
        if b in (0xD8, 0xD9) or b == 0x01 or 0xD0 <= b <= 0xD7:
            segs.append(b); i += 2; continue
        if b == 0xDA:
            break
        if i + 4 > n:
            break
        seglen = (data[i + 2] << 8) | data[i + 3]
        segs.append(b)
        i += 2 + seglen
    return segs


def make_png(path):
    from PIL.PngImagePlugin import PngInfo
    img = Image.new("RGB", (32, 32), (10, 120, 200))
    info = PngInfo()
    info.add_text("Author", "Alice")
    info.add_text("Software", "LeakyEditor")
    img.save(path, "PNG", pnginfo=info)


def make_webp(path):
    img = Image.new("RGB", (32, 32), (30, 200, 90))
    exif = Image.Exif()
    exif[0x0110] = "WebCam3000"
    img.save(path, "WEBP", exif=exif.tobytes())


def make_pdf(path):
    import pikepdf
    pdf = pikepdf.new()
    pdf.add_blank_page(page_size=(200, 200))
    with pdf.open_metadata() as m:
        m["dc:title"] = "Secret Report"
        m["dc:creator"] = ["Alice"]
    pdf.docinfo["/Title"] = "Secret Report"
    pdf.docinfo["/Author"] = "Alice"
    pdf.docinfo["/Producer"] = "LeakySoft"
    # attach an embedded file
    notes = os.path.join(TMP, "_att_notes.txt")
    open(notes, "w").write("top secret notes")
    pdf.attachments["notes.txt"] = pikepdf.AttachedFileSpec.from_filepath(pdf, notes)
    pdf.save(path)


def make_docx(path):
    import docx
    d = docx.Document()
    d.core_properties.author = "Alice"
    d.core_properties.last_modified_by = "Bob"
    d.core_properties.title = "Top Secret"
    d.core_properties.comments = "internal only"
    d.add_paragraph("hello world")
    d.save(path)
    # inject a picture with EXIF into word/media
    media = io.BytesIO()
    make_jpeg(media)
    exif_bytes = media.getvalue()
    with zipfile.ZipFile(path, "a") as z:
        z.writestr("word/media/image1.jpeg", exif_bytes)
    return exif_bytes


def office_xml(path, part):
    with zipfile.ZipFile(path) as z:
        return z.read(part).decode("utf-8", "ignore")


def main():
    shutil.rmtree(TMP, ignore_errors=True)
    os.makedirs(OUT, exist_ok=True)
    opts = Options()

    print("\n[1] JPEG — EXIF/GPS/XMP/comment removal (lossless)")
    j = os.path.join(TMP, "photo.jpg")
    make_jpeg(j)
    before_segs = read_jpeg_segments(j)
    check("source JPEG has APP1/EXIF", 0xE1 in before_segs, str(before_segs))
    rep = sanitize_file(j, default_output_path(j, OUT, "_clean"), opts)
    check("status ok", rep.status == "ok", rep.summary())
    segs = read_jpeg_segments(rep.dst)
    check("no APP1 (EXIF/XMP) left", 0xE1 not in segs, str(segs))
    check("no COM comment left", 0xFE not in segs, str(segs))
    check("no APP13 Photoshop left", 0xED not in segs, str(segs))
    with Image.open(rep.dst) as im:
        check("EXIF dict empty", not dict(im.getexif()), str(dict(im.getexif())))
    # quality: pixel content preserved (same dimensions; close colors)
    with Image.open(j) as a, Image.open(rep.dst) as b:
        check("same size", a.size == b.size)
        pa, pb = a.convert("RGB").getpixel((0, 0)), b.convert("RGB").getpixel((0, 0))
        check("pixels preserved losslessly", pa == pb, f"{pa} vs {pb}")
    src_bytes = open(j, "rb").read()
    out_bytes = open(rep.dst, "rb").read()
    check("output smaller (metadata removed)", len(out_bytes) < len(src_bytes))

    print("\n[2] JPEG — already clean")
    j2 = os.path.join(TMP, "plain.jpg")
    Image.new("RGB", (20, 20), (0, 0, 0)).save(j2, "JPEG")
    rep = sanitize_file(j2, default_output_path(j2, OUT, "_clean"), opts)
    check("status clean", rep.status == "clean", rep.summary())
    check("output identical bytes", open(rep.dst, "rb").read()
          == open(j2, "rb").read())

    print("\n[3] PNG — text chunk removal")
    p = os.path.join(TMP, "img.png")
    make_png(p)
    rep = sanitize_file(p, default_output_path(p, OUT, "_clean"), opts)
    check("status ok", rep.status == "ok", rep.summary())
    out = open(rep.dst, "rb").read()
    i, n = 8, len(out)
    chunks = []
    while i + 8 <= n:
        length = struct.unpack(">I", out[i:i + 4])[0]
        chunks.append(out[i + 4:i + 8])
        i += 8 + length + 4
    check("no tEXt chunks", b"tEXt" not in chunks, str(chunks))
    with Image.open(rep.dst) as im:
        check("pixel preserved", im.convert("RGB").getpixel((0, 0))
              == (10, 120, 200))

    print("\n[4] WebP — EXIF removal")
    w = os.path.join(TMP, "img.webp")
    make_webp(w)
    rep = sanitize_file(w, default_output_path(w, OUT, "_clean"), opts)
    check("status ok", rep.status == "ok", rep.summary())
    out = open(rep.dst, "rb").read()
    check("no EXIF chunk", b"EXIF" not in out)
    with Image.open(w) as a, Image.open(rep.dst) as b:
        check("pixel preserved", a.convert("RGB").getpixel((0, 0))
              == b.convert("RGB").getpixel((0, 0)))
        check("same size", a.size == b.size)

    print("\n[5] PDF — docinfo/XMP/attachments removal")
    pdf = os.path.join(TMP, "doc.pdf")
    make_pdf(pdf)
    rep = sanitize_file(pdf, default_output_path(pdf, OUT, "_clean"), opts)
    check("status ok", rep.status == "ok", rep.summary())
    import pikepdf
    with pikepdf.open(rep.dst) as q:
        check("docinfo empty", len(list(q.docinfo.keys())) == 0,
              str(list(q.docinfo.keys())))
        check("no XMP", "/Metadata" not in q.Root)
        check("no attachments", len(list(q.attachments.keys())) == 0,
              str(list(q.attachments.keys())))

    print("\n[6] DOCX — author & media EXIF removal")
    d = os.path.join(TMP, "report.docx")
    exif_bytes = make_docx(d)
    core_before = office_xml(d, "docProps/core.xml")
    check("source has author", "Alice" in core_before)
    rep = sanitize_file(d, default_output_path(d, OUT, "_clean"), opts)
    check("status ok", rep.status == "ok", rep.summary())
    core = office_xml(rep.dst, "docProps/core.xml")
    check("author removed", "Alice" not in core and "Bob" not in core)
    media = None
    with zipfile.ZipFile(rep.dst) as z:
        media = z.read("word/media/image1.jpeg")
    _, removed = strip_jpeg_bytes(media, keep_icc=True)
    check("embedded image has no EXIF", b"Exif\x00\x00" not in media)
    # doc still opens
    import docx
    doc = docx.Document(rep.dst)
    check("document still valid & has text",
          "hello world" in [p.text for p in doc.paragraphs])

    print("\n[7] Unsupported / missing")
    txt = os.path.join(TMP, "note.txt")
    open(txt, "w").write("hi")
    rep = sanitize_file(txt, default_output_path(txt, OUT, "_clean"), opts)
    check("skipped", rep.status == "skipped", rep.summary())
    rep = sanitize_file(os.path.join(TMP, "nope.jpg"),
                        default_output_path(txt, OUT, "_clean"), opts)
    check("missing -> error", rep.status == "error", rep.summary())

    print(f"\n===== {PASS} passed, {FAIL} failed =====")
    sys.exit(1 if FAIL else 0)


if __name__ == "__main__":
    main()
