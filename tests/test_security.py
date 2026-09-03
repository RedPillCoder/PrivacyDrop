"""Security-hardening verification suite for the PrivacyDrop engine.

Covers: decompression-bomb / zip-bomb defence, zip-slip / path-traversal /
absolute-path / drive-letter entry rejection, encrypted zip handling,
malformed zips, XML entity immunity (no XML parsing of untrusted data),
VBA macro stripping, PDF JavaScript / launch-action / XFA stripping (while
preserving benign GoTo links), atomic writes (no stray temp files, clean
overwrite), input-size limits, and no-XML-parser guarantees.

Run:  python tests/test_security.py
"""
import os
import shutil
import struct
import sys
import zipfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import scrubber
from scrubber import (sanitize_file, Options, default_output_path, ScrubError,
                      _read_zip_safely, _check_zip_entry_name,
                      _strip_pdf_unsafe)

from PIL import Image
from PIL.TiffImagePlugin import IFDRational

TMP = os.path.join(os.path.dirname(os.path.abspath(__file__)), "tmp_sec")
OUT = os.path.join(TMP, "out")
PASS = FAIL = 0


def check(name, cond, detail=""):
    global PASS, FAIL
    if cond:
        PASS += 1
        print(f"  PASS  {name}")
    else:
        FAIL += 1
        print(f"  FAIL  {name}   {detail}")


def sanitize(src, opts=None, out_dir=OUT):
    opts = opts or Options()
    dst = default_output_path(src, out_dir, opts.suffix)
    return sanitize_file(src, dst, opts)


def make_exif():
    exif = Image.Exif()
    exif[0x010F] = "Acme Corp"
    exif[0x0110] = "SecretCam 9000"
    exif[0x8825] = {1: "S",
                    2: (IFDRational(31, 1), IFDRational(57, 1), IFDRational(1, 1)),
                    3: "E",
                    4: (IFDRational(115, 1), IFDRational(51, 1), IFDRational(1, 1))}
    return exif


def make_docx(path, extra_entries=None, core_override=None):
    import docx
    d = docx.Document()
    d.core_properties.author = "Alice"
    d.add_paragraph("hello")
    d.save(path)
    if extra_entries or core_override is not None:
        with zipfile.ZipFile(path, "a") as z:
            if core_override is not None:
                z.writestr("docProps/core.xml", core_override)
            for name, data in (extra_entries or {}).items():
                z.writestr(name, data)


def make_flagged_encrypted_zip(path, entries):
    """Build a zip whose entries carry the encryption flag (simulating a
    password-protected Office package; stdlib zipfile can't create real
    encryption, but this exercises the exact RuntimeError path we defend)."""
    with zipfile.ZipFile(path, "w") as z:
        for n, d in entries.items():
            z.writestr(n, d)
    data = bytearray(open(path, "rb").read())
    for sig, off in ((b"PK\x03\x04", 6), (b"PK\x01\x02", 8)):
        i = 0
        while True:
            j = data.find(sig, i)
            if j == -1:
                break
            data[j + off] |= 0x01
            i = j + 4
    open(path, "wb").write(data)


# ---------------------------------------------------------------- zip-slip

def t_zip_entry_names():
    print("\n[Zip entry-name defence]")
    cases = [
        ("../evil.txt", "path-traversal"),
        ("a/../../evil.txt", "path-traversal"),
        ("/etc/passwd", "absolute-path"),
        ("\\windows\\system32\\evil.dll", "absolute-path (backslash)"),
        ("C:\\evil.exe", "drive-letter"),
        ("C:/evil.exe", "drive-letter (fwd slash)"),
        ("a\x00b.xml", "null byte"),
        ("", "empty"),
    ]
    for name, why in cases:
        try:
            _check_zip_entry_name(name)
            check(f"reject {why}: {name!r}", False, "no exception raised")
        except ScrubError:
            check(f"reject {why}: {name!r}", True)
    # benign names accepted
    for name in ("[Content_Types].xml", "word/document.xml",
                 "docProps/core.xml", "word/media/image1.png"):
        try:
            _check_zip_entry_name(name)
            check(f"accept benign: {name!r}", True)
        except ScrubError as e:
            check(f"accept benign: {name!r}", False, str(e))


def t_zip_bomb():
    print("\n[Decompression-bomb defence]")
    # a zip whose declared uncompressed size explodes
    bomb = os.path.join(TMP, "bomb.docx")
    blob = b"A" * (64 * 1024 * 1024)          # 64 MiB of zeros -> tiny zip
    with zipfile.ZipFile(bomb, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("word/document.xml", blob)
        z.writestr("docProps/core.xml", scrubber.CORE_XML)
    size_on_disk = os.path.getsize(bomb)
    rep = sanitize(bomb)
    # with the default 512 MiB/entry cap this is *not* a bomb — it should
    # actually succeed; this guards the "not a bomb" path.
    check("large-but-legal zip processed", rep.ok, rep.summary())

    # a true bomb: exceed per-entry cap by monkeypatching the limit
    old = scrubber.MAX_ZIP_ENTRY_UNCOMPRESSED
    scrubber.MAX_ZIP_ENTRY_UNCOMPRESSED = 1024 * 1024  # 1 MiB
    try:
        try:
            _read_zip_safely(bomb)
            check("oversized entry rejected", False, "no exception")
        except ScrubError as e:
            check("oversized entry rejected", "too large" in str(e), str(e))
    finally:
        scrubber.MAX_ZIP_ENTRY_UNCOMPRESSED = old

    # entry count cap
    oldc = scrubber.MAX_ZIP_ENTRIES
    scrubber.MAX_ZIP_ENTRIES = 5
    many = os.path.join(TMP, "many.docx")
    with zipfile.ZipFile(many, "w") as z:
        for i in range(20):
            z.writestr(f"word/part{i}.xml", "<w/>")
    try:
        try:
            _read_zip_safely(many)
            check("too-many-entries rejected", False, "no exception")
        except ScrubError as e:
            check("too-many-entries rejected", "too many entries" in str(e), str(e))
    finally:
        scrubber.MAX_ZIP_ENTRIES = oldc

    # path-traversal entry inside a real zip -> rejected end-to-end
    trav = os.path.join(TMP, "trav.docx")
    with zipfile.ZipFile(trav, "w") as z:
        z.writestr("word/document.xml", "<w/>")
        z.writestr("../../evil.txt", "pwn")
    rep = sanitize(trav)
    check("traversal zip -> error", rep.status == "error"
          and "path-traversal" in rep.error, rep.error)
    check("no output written", not os.path.exists(rep.dst))


def t_encrypted_malformed():
    print("\n[Encrypted & malformed archives]")
    # password-protected (encrypted) Office package
    enc = os.path.join(TMP, "enc.docx")
    make_flagged_encrypted_zip(enc, {
        "word/document.xml": "<w/>",
        "docProps/core.xml": scrubber.CORE_XML.encode(),
    })
    rep = sanitize(enc)
    check("encrypted zip -> clear error",
          rep.status == "error" and "encrypted" in rep.error.lower(),
          rep.error)
    check("encrypted zip -> no output", not os.path.exists(rep.dst))

    # not a zip at all
    bad = os.path.join(TMP, "fake.docx")
    open(bad, "wb").write(b"this is not a zip file at all")
    rep = sanitize(bad)
    check("non-zip docx -> error", rep.status == "error", rep.summary())

    # empty file
    empty = os.path.join(TMP, "empty.docx")
    open(empty, "wb").write(b"")
    rep = sanitize(empty)
    check("empty docx -> error", rep.status == "error", rep.summary())


def t_xml_entity_immunity():
    print("\n[XML entity immunity (no XML parsing of untrusted data)]")
    # core.xml with a DOCTYPE + external entity + billion-laughs style entity
    evil_core = (
        '<?xml version="1.0"?>'
        '<!DOCTYPE lolz ['
        '<!ENTITY lol "lol"><!ENTITY lol2 "&lol;&lol;&lol;&lol;">'
        '<!ENTITY xxe SYSTEM "file:///etc/passwd">]>'
        '<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/'
        'package/2006/metadata/core-properties"'
        ' xmlns:dc="http://purl.org/dc/elements/1.1/">'
        '<dc:creator>&xxe;Alice</dc:creator>'
        '<cp:lastModifiedBy>Bob</cp:lastModifiedBy>'
        '</cp:coreProperties>'
    )
    d = os.path.join(TMP, "evil.docx")
    make_docx(d, core_override=evil_core.encode())
    rep = sanitize(d)
    check("entity-laden docx processed without crash", rep.status == "ok",
          rep.summary())
    core = None
    with zipfile.ZipFile(rep.dst) as z:
        core = z.read("docProps/core.xml").decode()
    check("output core.xml is the safe template", "&xxe;" not in core
          and "<!DOCTYPE" not in core and "Alice" not in core and
          "creator></dc:creator>" in core)
    check("report still lists removed fields", "creator" in rep.summary()
          or "document properties" in rep.summary(), rep.summary())


def t_macro_stripping():
    print("\n[VBA macro stripping]")
    d = os.path.join(TMP, "macro.docx")
    # add a real JPEG to media for good measure
    img = os.path.join(TMP, "_m.jpg")
    Image.new("RGB", (20, 20), (1, 2, 3)).save(img, "JPEG",
                                               exif=make_exif().tobytes())
    make_docx(d, extra_entries={
        "word/vbaProject.bin": b"\xd0\xcf\x11\xe0" + b"M" * 100,   # OLE header
        "word/vbaProjectSignature.bin": b"sig" * 20,
        "word/media/image1.jpeg": open(img, "rb").read(),
    })

    rep = sanitize(d)   # remove_scripts=True by default
    check("macro docx cleaned", rep.status == "ok", rep.summary())
    check("macro removal reported", "macro" in rep.summary().lower(),
          rep.summary())
    with zipfile.ZipFile(rep.dst) as z:
        names = z.namelist()
        check("vbaProject.bin removed", "word/vbaProject.bin" not in names)
        check("vbaProjectSignature removed",
              "word/vbaProjectSignature.bin" not in names)
        check("media preserved", "word/media/image1.jpeg" in names)
        media = z.read("word/media/image1.jpeg")
        check("media still scrubbed", b"Exif\x00\x00" not in media)
    import docx
    doc = docx.Document(rep.dst)
    check("docx still opens", "hello" in [p.text for p in doc.paragraphs])

    # with remove_scripts=False macros are kept
    rep2 = sanitize(d, Options(remove_scripts=False))
    with zipfile.ZipFile(rep2.dst) as z:
        check("macros kept when opted out",
              "word/vbaProject.bin" in z.namelist())


def make_js_pdf(path):
    import pikepdf
    pdf = pikepdf.new()
    pdf.add_blank_page(page_size=(200, 200))
    pdf.docinfo["/Author"] = "Alice"
    # annotation JS
    annot = pikepdf.Dictionary({
        "/Type": pikepdf.Name("/Annot"),
        "/Subtype": pikepdf.Name("/Widget"),
        "/Rect": pikepdf.Array([0, 0, 10, 10]),
        "/JS": pikepdf.String("this.print(true)"),
    })
    pdf.pages[0]["/Annots"] = pikepdf.Array([annot])
    # additional actions: one JS (remove), one GoTo (keep)
    pdf.pages[0]["/AA"] = pikepdf.Dictionary({
        "/O": pikepdf.Dictionary({"/S": pikepdf.Name("/JavaScript"),
                                  "/JS": pikepdf.String("run()")}),
        "/C": pikepdf.Dictionary({"/S": pikepdf.Name("/GoTo"),
                                  "/D": pikepdf.String("dest")}),
    })
    # open action JS
    pdf.Root["/OpenAction"] = pikepdf.Dictionary(
        {"/S": pikepdf.Name("/JavaScript"),
         "/JS": pikepdf.String("app.launchURL('http://evil')")})
    # document-level JS tree
    pdf.Root["/Names"] = pikepdf.Dictionary({
        "/JavaScript": pikepdf.Dictionary({
            "/Names": pikepdf.Array([pikepdf.String("x"),
                                     pikepdf.Dictionary({"/S": pikepdf.Name("/JavaScript"),
                                                         "/JS": pikepdf.String("boom()")})]),
        }),
    })
    # launch action (executes a program)
    pdf.pages[0]["/A"] = pikepdf.Dictionary(
        {"/S": pikepdf.Name("/Launch"), "/F": pikepdf.String("calc.exe")})
    pdf.save(path)


def t_pdf_scripts():
    print("\n[PDF JavaScript / launch-action / XFA stripping]")
    import pikepdf

    def collect_js_keys(pdf):
        found = []

        def scan(obj, depth):
            if depth > 64:
                return
            try:
                keys = list(obj.keys())
            except Exception:
                try:
                    for i in range(len(obj)):
                        scan(obj[i], depth + 1)
                except Exception:
                    pass
                return
            for k in keys:
                if str(k) in ("/JS", "/JavaScript"):
                    found.append(str(k))
                scan(obj[k], depth + 1)
        for o in pdf.objects:
            scan(o, 0)
        return found

    p = os.path.join(TMP, "js.pdf")
    make_js_pdf(p)
    src = pikepdf.open(p)
    check("fixture has JS", bool(collect_js_keys(src)))
    src.close()

    rep = sanitize(p)          # remove_scripts=True by default
    check("js pdf cleaned", rep.status == "ok", rep.summary())
    check("scripts reported", "script" in rep.summary(), rep.summary())

    out = pikepdf.open(rep.dst)
    check("no /JS or /JavaScript keys remain", not collect_js_keys(out))
    page = out.pages[0]
    pk = list(page.obj.keys())
    check("open action removed", "/OpenAction" not in list(out.Root.keys()))
    check("launch action removed", "/A" not in pk)
    # benign GoTo action inside /AA preserved
    if "/AA" in pk:
        aa = list(page.obj["/AA"].keys())
        check("GoTo action preserved", "/C" in aa and "/O" not in aa, str(aa))
    check("author still removed", not list(out.docinfo.keys()))
    out.close()

    # opt-out keeps scripts
    rep2 = sanitize(p, Options(remove_scripts=False))
    out2 = pikepdf.open(rep2.dst)
    check("scripts kept when opted out", bool(collect_js_keys(out2)))
    out2.close()


def t_atomic_write():
    print("\n[Atomic write behaviour]")
    # no stray temp files after a successful run
    src = os.path.join(TMP, "atom.jpg")
    Image.new("RGB", (20, 20), (5, 5, 5)).save(src, "JPEG",
                                               exif=make_exif().tobytes())
    rep = sanitize(src)
    check("cleaned ok", rep.status == "ok")
    leftovers = [f for f in os.listdir(OUT)
                 if ".part" in f or f.startswith(".privacydrop-")]
    check("no stray temp files", not leftovers, str(leftovers))

    # overwrite: run twice, second run must bump name (collision) and succeed
    rep2 = sanitize(src)
    check("second run bumps name", rep2.dst != rep.dst and rep2.ok)
    check("both outputs exist",
          os.path.isfile(rep.dst) and os.path.isfile(rep2.dst))

    # failed write must not leave a temp file: simulate by writing into a
    # read-only directory (skip if running as root, chmod won't stop root)
    if os.geteuid() != 0:
        rodir = os.path.join(TMP, "ro")
        os.makedirs(rodir)
        os.chmod(rodir, 0o555)
        dst = os.path.join(rodir, "x_clean.jpg")
        try:
            from scrubber import _write
            try:
                _write(b"data", dst)
                check("write to read-only dir fails", False, "should have raised")
            except OSError:
                check("write to read-only dir fails", True)
        finally:
            os.chmod(rodir, 0o755)
        check("no temp file left after failure",
              not [f for f in os.listdir(rodir) if f != "x_clean.jpg"])


def t_input_size_limit():
    print("\n[Input size limits]")
    old = scrubber.MAX_INPUT_BYTES
    scrubber.MAX_INPUT_BYTES = 1024     # tiny cap to make the fixture "oversized"
    try:
        big = os.path.join(TMP, "big.jpg")
        # random noise JPEG — incompressible, guaranteed > 1 KB
        import random
        random.seed(7)
        noise = Image.new("RGB", (200, 200))
        noise.putdata([(random.randrange(256), random.randrange(256),
                        random.randrange(256)) for _ in range(200 * 200)])
        noise.save(big, "JPEG", exif=make_exif().tobytes())
        rep = sanitize(big)
        check("oversized image rejected with clear error",
              rep.status == "error" and "too large" in rep.error,
              rep.error)
        check("no output written", not os.path.exists(rep.dst))
    finally:
        scrubber.MAX_INPUT_BYTES = old


def main():
    shutil.rmtree(TMP, ignore_errors=True)
    os.makedirs(OUT)
    t_zip_entry_names()
    t_zip_bomb()
    t_encrypted_malformed()
    t_xml_entity_immunity()
    t_macro_stripping()
    t_pdf_scripts()
    t_atomic_write()
    t_input_size_limit()
    print(f"\n===== {PASS} passed, {FAIL} failed =====")
    sys.exit(1 if FAIL else 0)


if __name__ == "__main__":
    main()
